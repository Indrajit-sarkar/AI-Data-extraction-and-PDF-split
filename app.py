from flask import Flask, request, jsonify, send_from_directory, send_file
from flask_cors import CORS
import os
import sys
import io
import logging
import traceback
from functools import wraps

# Fix Windows console encoding for Unicode characters
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

import json
import time
import shutil

from pathlib import Path
from datetime import datetime
import tempfile
import importlib
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# ============================================================================
# CONFIGURATION MANAGEMENT
# ============================================================================
class Config:
    """Centralized configuration from environment variables"""
    # Server
    SERVER_PORT = int(os.getenv('SERVER_PORT', 5000))
    SERVER_HOST = os.getenv('SERVER_HOST', '0.0.0.0')
    DEBUG_MODE = os.getenv('DEBUG_MODE', 'true').lower() == 'true'
    
    # Processing
    MAX_THREADS = int(os.getenv('MAX_THREADS', 4))
    MAX_PROCESSES = int(os.getenv('MAX_PROCESSES', 2))
    RETRY_ATTEMPTS = int(os.getenv('RETRY_ATTEMPTS', 3))
    RETRY_DELAY_BASE = float(os.getenv('RETRY_DELAY_BASE', 2.0))
    CONTENT_TRUNCATE_LIMIT = int(os.getenv('CONTENT_TRUNCATE_LIMIT', 50000))
    
    # Rate Limiting
    RATE_LIMIT_DEFAULT = os.getenv('RATE_LIMIT_DEFAULT', '100 per minute')
    RATE_LIMIT_PROCESSING = os.getenv('RATE_LIMIT_PROCESSING', '20 per minute')
    RATE_LIMIT_ENABLED = os.getenv('RATE_LIMIT_ENABLED', 'true').lower() == 'true'
    
    # CORS
    CORS_ORIGINS = os.getenv('CORS_ORIGINS', '*')
    
    # Logging
    LOG_LEVEL = os.getenv('LOG_LEVEL', 'INFO')
    LOG_FORMAT = os.getenv('LOG_FORMAT', 'json')
    
    # PDF Splitting Configuration
    SPLIT_PDF_OUTPUT_PATH = os.getenv('SPLIT_PDF_OUTPUT_PATH', 
        os.path.join(os.path.dirname(__file__), 'SPLIT PDF FILES'))
    SPLIT_PDF_ENABLED = os.getenv('SPLIT_PDF_ENABLED', 'true').lower() == 'true'
    SPLIT_PDF_METHOD = os.getenv('SPLIT_PDF_METHOD', 'barcode_vendor')
    SPLIT_PDF_NAMING_PATTERN = os.getenv('SPLIT_PDF_NAMING_PATTERN', '{original_name}_{increment}')
    SPLIT_PDF_AUTO_PROCESS = os.getenv('SPLIT_PDF_AUTO_PROCESS', 'true').lower() == 'true'
    SPLIT_PDF_AUTO_EXTRACT = os.getenv('SPLIT_PDF_AUTO_EXTRACT', 'true').lower() == 'true'
    

    
    @classmethod
    def validate(cls):
        """Validate required configuration"""
        required_vars = [
            ('DOC_ENDPOINT', os.getenv('DOC_ENDPOINT')),
            ('DOC_KEY', os.getenv('DOC_KEY')),
        ]
        missing = [name for name, value in required_vars if not value]
        if missing:
            logger.warning(f"Missing recommended env vars: {missing}")
        return True

# ============================================================================
# STRUCTURED LOGGING
# ============================================================================
class StructuredFormatter(logging.Formatter):
    """JSON structured log formatter"""
    def format(self, record):
        log_data = {
            'timestamp': datetime.now().astimezone().isoformat(),
            'level': record.levelname,
            'message': record.getMessage(),
            'module': record.module,
            'function': record.funcName,
            'line': record.lineno
        }
        if hasattr(record, 'request_id'):
            log_data['request_id'] = record.request_id
        if record.exc_info:
            log_data['exception'] = self.formatException(record.exc_info)
        return json.dumps(log_data, ensure_ascii=False)

# Configure logging
log_level = getattr(logging, Config.LOG_LEVEL.upper(), logging.INFO)
logger = logging.getLogger('universal_extractor')
logger.setLevel(log_level)

handler = logging.StreamHandler()
if Config.LOG_FORMAT == 'json':
    handler.setFormatter(StructuredFormatter())
else:
    handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))
logger.addHandler(handler)

# ============================================================================
# CUSTOM EXCEPTIONS
# ============================================================================
class ExtractorException(Exception):
    """Base exception for extractor errors"""
    def __init__(self, message, status_code=500, details=None):
        super().__init__(message)
        self.message = message
        self.status_code = status_code
        self.details = details or {}

class ValidationError(ExtractorException):
    """Input validation error"""
    def __init__(self, message, details=None):
        super().__init__(message, status_code=400, details=details)

class ProcessingError(ExtractorException):
    """Document processing error"""
    def __init__(self, message, details=None):
        super().__init__(message, status_code=500, details=details)

class NotFoundError(ExtractorException):
    """Resource not found error"""
    def __init__(self, message, details=None):
        super().__init__(message, status_code=404, details=details)

# ============================================================================
# PATH VALIDATION (Security)
# ============================================================================
def validate_path(path: str, allowed_roots: list = None) -> Path:
    """
    Validate and sanitize file path to prevent path traversal attacks.
    Returns resolved absolute path or raises ValidationError.
    """
    try:
        # Resolve to absolute path
        resolved = Path(path).resolve()
        
        # Check for path traversal attempts
        if '..' in path:
            logger.warning(f"Path traversal attempt detected: {path}")
            raise ValidationError("Invalid path: path traversal not allowed")
        
        # Optionally check against allowed roots
        if allowed_roots:
            is_allowed = any(
                str(resolved).startswith(str(Path(root).resolve())) 
                for root in allowed_roots
            )
            if not is_allowed:
                logger.warning(f"Path outside allowed roots: {path}")
                raise ValidationError("Path outside allowed directories")
        
        return resolved
    except ValidationError:
        raise
    except Exception as e:
        raise ValidationError(f"Invalid path: {str(e)}")


# Initialize global in-memory session
review_session = {
    'documents': [],
    'extracted_data': {},
    'json_files': {},
    'current_index': 0
}
# ============================================================================


# Initialize config and validate
Config.validate()
logger.info("Configuration loaded and validated")

# Import the extractor modules
sys.path.append(os.path.dirname(__file__))

extractor = None
pdf_processor = None
parallel_processor = None
pdf_splitter = None

# Try to import multi_file_extractor
possible_names = ['multi_file_extractor']
for module_name in possible_names:
    try:
        extractor = importlib.import_module(module_name)
        print(f"[OK] Successfully imported: {module_name}")
        break
    except ImportError:
        continue

# Try to import Barcode & Vendor PDF Splitter (pdf.py - focused on barcode/vendor detection)
try:
    from pdf import BarcodeVendorPDFSplitter
    pdf_processor_instance = BarcodeVendorPDFSplitter()
    print(f"[OK] Successfully imported: pdf.py (Barcode & Vendor Splitter)")
    print(f"[OK] PDF output directory: {Config.SPLIT_PDF_OUTPUT_PATH}")
    
    # Check dependencies
    deps = pdf_processor_instance.check_dependencies()
    if deps['missing']:
        print(f"[!!] Optional libraries missing: {', '.join(deps['missing'])}")
        print(f"[!!] Install with: pip install {' '.join(deps['missing'])}")
    
    if deps['features'].get('barcode_detection', False):
        print(f"[OK] Barcode Detection: ENABLED (Priority #1)")
    else:
        print(f"[!!] Barcode Detection: DISABLED - Install: pip install pdf2image pyzbar pillow")
    
    print(f"[OK] Vendor Extraction: ENABLED (Priority #2)")
    
    # Keep backward compatibility references
    pdf_processor = pdf_processor_instance
    pdf_splitter = pdf_processor_instance
except ImportError as e:
    print(f"[!!] Warning: pdf.py not found - {e}")
    pdf_processor = None
    pdf_splitter = None
except Exception as e:
    print(f"[!!] Warning: PDF processor initialization failed - {e}")
    import traceback
    print(traceback.format_exc())
    pdf_processor = None
    pdf_splitter = None

# Try to import parallel processor (primary processor)
try:
    parallel_processor = importlib.import_module('parallel_document_processor')
    print(f"[OK] Successfully imported: parallel_document_processor.py")
except ImportError as e:
    print(f"[!!] Warning: parallel_document_processor.py not found - {e}")

if extractor is None and parallel_processor is None:
    print("=" * 80)
    print("[ERROR] Could not import any extractor module")
    print("=" * 80)

# ============================================================================
# DOCUMENT REVIEW CONFIGURATION
# ============================================================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# JSON output folders
JSON_FILES_FOLDER = os.getenv('OUTPUT_JSON_FILES_PATH', 
    os.path.join(BASE_DIR, 'OUTPUT JSON', 'JSON FILES'))
ACCEPTED_FOLDER = os.getenv('OUTPUT_JSON_ACCEPTED_PATH', 
    os.path.join(BASE_DIR, 'OUTPUT JSON', 'ACCEPTED'))
REJECTED_FOLDER = os.getenv('OUTPUT_JSON_REJECTED_PATH', 
    os.path.join(BASE_DIR, 'OUTPUT JSON', 'REJECTED'))

# PDF Split output folder
SPLIT_PDF_FOLDER = Config.SPLIT_PDF_OUTPUT_PATH

# Ensure folders exist
os.makedirs(JSON_FILES_FOLDER, exist_ok=True)
os.makedirs(ACCEPTED_FOLDER, exist_ok=True)
os.makedirs(REJECTED_FOLDER, exist_ok=True)
os.makedirs(SPLIT_PDF_FOLDER, exist_ok=True)

logger.info(f"JSON output folders configured: {JSON_FILES_FOLDER}")
logger.info(f"PDF split folder configured: {SPLIT_PDF_FOLDER}")



def get_file_type(filename):
    """Detect file type from extension"""
    ext = Path(filename).suffix.lower()
    type_map = {
        '.pdf': 'PDF',
        '.docx': 'DOCX',
        '.doc': 'DOCX',
        '.eml': 'EML',
        '.xlsx': 'EXCEL',
        '.xls': 'EXCEL',
        '.csv': 'CSV'
    }
    return type_map.get(ext, 'UNKNOWN')

logger.info(f"Document Review paths configured - Accepted: {ACCEPTED_FOLDER}, Rejected: {REJECTED_FOLDER}")

# ============================================================================
# FLASK APP WITH CONFIGURED CORS
# ============================================================================
app = Flask(__name__)

# Configure CORS with specific origins
cors_origins = Config.CORS_ORIGINS
if cors_origins != '*':
    cors_origins = [origin.strip() for origin in cors_origins.split(',')]
CORS(app, origins=cors_origins)

# ============================================================================
# ERROR HANDLERS
# ============================================================================
@app.errorhandler(ExtractorException)
def handle_extractor_error(error):
    """Handle custom extractor exceptions"""
    logger.error(f"ExtractorException: {error.message}", exc_info=True)
    return jsonify({
        'success': False,
        'error': error.message,
        'details': error.details
    }), error.status_code

@app.errorhandler(Exception)
def handle_generic_error(error):
    """Handle unexpected exceptions"""
    logger.error(f"Unexpected error: {str(error)}", exc_info=True)
    return jsonify({
        'success': False,
        'error': 'An unexpected error occurred',
        'details': {'type': type(error).__name__}
    }), 500

# ============================================================================
# HEALTH CHECK ENDPOINTS
# ============================================================================
@app.route('/health', methods=['GET'])
def health_check():
    """Basic health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.utcnow().isoformat() + 'Z',
        'version': '2.0.0'
    })

@app.route('/ready', methods=['GET'])
def readiness_check():
    """Readiness check - verifies all dependencies are available"""
    checks = {
        'parallel_processor': parallel_processor is not None,
        'extractor': extractor is not None,
        'json_folder': os.path.exists(JSON_FILES_FOLDER),
        'accepted_folder': os.path.exists(ACCEPTED_FOLDER),
        'rejected_folder': os.path.exists(REJECTED_FOLDER)
    }
    
    all_ready = all(checks.values())
    
    return jsonify({
        'status': 'ready' if all_ready else 'not_ready',
        'timestamp': datetime.utcnow().isoformat() + 'Z',
        'checks': checks
    }), 200 if all_ready else 503

# ============================================================================
# MAIN APP ROUTES
# ============================================================================
@app.route('/')
def index():
    return send_from_directory('.', 'index.html')

@app.route('/review')
def document_review_page():
    """Serve the Document Review HTML application"""
    return send_from_directory('.', 'document_review.html')

@app.route('/test_upload.html')
def test_upload_page():
    """Serve the test upload page"""
    return send_from_directory('.', 'test_upload.html')

@app.route('/simple_upload.html')
def simple_upload_page():
    """Serve the simple upload page"""
    return send_from_directory('.', 'simple_upload.html')

@app.route('/pdf_split_test.html')
def pdf_split_test_page():
    """Serve the PDF split test page"""
    return send_from_directory('.', 'pdf_split_test.html')

@app.route('/bg.mp4')
def serve_video():
    return send_from_directory('.', 'bg.mp4')

@app.route('/favicon.ico')
def favicon():
    """Serve favicon or return 204 if not found"""
    try:
        return send_from_directory('.', 'favicon.ico')
    except:
        # Return empty response if favicon not found
        from flask import Response
        return Response(status=204)

# ============================================================================
# SETTINGS MANAGEMENT ENDPOINTS
# ============================================================================
@app.route('/settings')
def settings_page():
    """Serve the settings HTML page"""
    return send_from_directory('.', 'settings.html')

@app.route('/api/settings', methods=['GET'])
def get_settings():
    """Get current settings from .env file"""
    try:
        env_path = os.path.join(BASE_DIR, '.env')
        settings = {}
        
        if os.path.exists(env_path):
            with open(env_path, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if line and not line.startswith('#') and '=' in line:
                        key, _, value = line.partition('=')
                        key = key.strip()
                        value = value.strip()
                        # Remove quotes if present
                        if (value.startswith('"') and value.endswith('"')) or \
                           (value.startswith("'") and value.endswith("'")):
                            value = value[1:-1]
                        settings[key] = value
        
        return jsonify(settings)
    except Exception as e:
        logger.error(f"Failed to read settings: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/settings', methods=['POST'])
def update_settings():
    """Update settings in .env file"""
    try:
        new_settings = request.get_json()
        if not new_settings:
            return jsonify({'error': 'No settings provided'}), 400
        
        env_path = os.path.join(BASE_DIR, '.env')
        
        # Read existing .env file
        existing_lines = []
        existing_keys = set()
        if os.path.exists(env_path):
            with open(env_path, 'r', encoding='utf-8') as f:
                existing_lines = f.readlines()
        
        # Parse existing keys
        for i, line in enumerate(existing_lines):
            stripped = line.strip()
            if stripped and not stripped.startswith('#') and '=' in stripped:
                key = stripped.split('=')[0].strip()
                existing_keys.add(key)
        
        # Update existing lines with new values
        updated_lines = []
        for line in existing_lines:
            stripped = line.strip()
            if stripped and not stripped.startswith('#') and '=' in stripped:
                key = stripped.split('=')[0].strip()
                if key in new_settings:
                    value = new_settings[key]
                    # Escape special characters in value
                    if isinstance(value, str) and any(c in value for c in [' ', '"', "'"]):
                        updated_lines.append(f'{key}="{value}"\n')
                    else:
                        updated_lines.append(f'{key}={value}\n')
                else:
                    updated_lines.append(line)
            else:
                updated_lines.append(line)
        
        # Add new keys that don't exist
        for key, value in new_settings.items():
            if key not in existing_keys and not key.startswith('_'):
                if isinstance(value, str) and any(c in value for c in [' ', '"', "'"]):
                    updated_lines.append(f'{key}="{value}"\n')
                else:
                    updated_lines.append(f'{key}={value}\n')
        
        # Write updated .env file
        with open(env_path, 'w', encoding='utf-8') as f:
            f.writelines(updated_lines)
        
        logger.info(f"Settings updated: {list(new_settings.keys())}")
        return jsonify({'success': True, 'message': 'Settings saved successfully'})
    
    except Exception as e:
        logger.error(f"Failed to update settings: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/settings/verify/doc-intelligence', methods=['POST'])
def verify_doc_intelligence():
    """Verify Document Intelligence endpoint and key"""
    import requests as http_requests
    try:
        data = request.get_json()
        endpoint = data.get('endpoint', '').strip()
        key = data.get('key', '').strip()
        
        if not endpoint or not key:
            return jsonify({'valid': False, 'error': 'Endpoint and key are required'})
        
        # Remove trailing slash from endpoint
        endpoint = endpoint.rstrip('/')
        
        # List of API versions to try (newest first)
        api_versions = [
            ('documentintelligence', '2024-02-29-preview'),
            ('documentintelligence', '2023-10-31-preview'),
            ('formrecognizer', '2023-07-31'),
            ('formrecognizer', '2022-08-31'),
        ]
        
        headers = {
            'Ocp-Apim-Subscription-Key': key,
            'Content-Type': 'application/json'
        }
        
        last_error = None
        for service_path, api_version in api_versions:
            try:
                test_url = f"{endpoint}/{service_path}/documentModels?api-version={api_version}"
                response = http_requests.get(test_url, headers=headers, timeout=15)
                
                if response.status_code == 200:
                    return jsonify({'valid': True, 'message': f'Connection successful (API: {api_version})'})
                elif response.status_code == 401:
                    return jsonify({'valid': False, 'error': 'Invalid API key'})
                elif response.status_code == 403:
                    return jsonify({'valid': False, 'error': 'Access forbidden - check API key permissions'})
                else:
                    last_error = f'API returned status {response.status_code}'
            except http_requests.exceptions.Timeout:
                last_error = 'Connection timeout'
            except http_requests.exceptions.ConnectionError:
                last_error = 'Could not connect to endpoint'
            except Exception as e:
                last_error = str(e)
        
        return jsonify({'valid': False, 'error': last_error or 'Could not verify endpoint'})
    
    except Exception as e:
        logger.error(f"Doc Intelligence verification failed: {e}")
        return jsonify({'valid': False, 'error': str(e)})

@app.route('/api/settings/verify/azure-openai', methods=['POST'])
def verify_azure_openai():
    """Verify Azure OpenAI endpoint and key"""
    import requests as http_requests
    import re
    try:
        data = request.get_json()
        endpoint = data.get('endpoint', '').strip()
        key = data.get('key', '').strip()
        api_version = data.get('apiVersion', '2024-08-01-preview').strip()
        deployment = data.get('deployment', '').strip()
        
        if not endpoint or not key:
            return jsonify({'valid': False, 'error': 'Endpoint and key are required'})
        
        # Extract base endpoint if full URL with deployment is provided
        # Handles: https://xxx.openai.azure.com/openai/deployments/model1/chat/completions?api-version=xxx
        base_url_match = re.match(r'(https://[^/]+\.openai\.azure\.com)', endpoint)
        if base_url_match:
            base_endpoint = base_url_match.group(1)
        else:
            base_endpoint = endpoint.rstrip('/')
        
        # Try to extract deployment from URL if not provided
        if not deployment:
            deploy_match = re.search(r'/deployments/([^/]+)', endpoint)
            if deploy_match:
                deployment = deploy_match.group(1)
        
        headers = {
            'api-key': key,
            'Content-Type': 'application/json'
        }
        
        # Try to list deployments first
        test_url = f"{base_endpoint}/openai/deployments?api-version={api_version}"
        try:
            response = http_requests.get(test_url, headers=headers, timeout=15)
            
            if response.status_code == 200:
                return jsonify({'valid': True, 'message': 'Connection successful'})
            elif response.status_code == 401:
                return jsonify({'valid': False, 'error': 'Invalid API key'})
        except:
            pass
        
        # Try with specific deployment
        if deployment:
            test_url = f"{base_endpoint}/openai/deployments/{deployment}?api-version={api_version}"
            try:
                response = http_requests.get(test_url, headers=headers, timeout=15)
                if response.status_code == 200:
                    return jsonify({'valid': True, 'message': f'Connection successful (deployment: {deployment})'})
                elif response.status_code == 401:
                    return jsonify({'valid': False, 'error': 'Invalid API key'})
            except:
                pass
        
        # Try a simple chat completion test
        if deployment:
            test_url = f"{base_endpoint}/openai/deployments/{deployment}/chat/completions?api-version={api_version}"
            try:
                test_body = {
                    "messages": [{"role": "user", "content": "test"}],
                    "max_tokens": 1
                }
                response = http_requests.post(test_url, headers=headers, json=test_body, timeout=15)
                if response.status_code in [200, 400]:  # 400 means endpoint works but request was malformed
                    return jsonify({'valid': True, 'message': f'Connection successful (deployment: {deployment})'})
                elif response.status_code == 401:
                    return jsonify({'valid': False, 'error': 'Invalid API key'})
            except:
                pass
        
        return jsonify({'valid': False, 'error': 'Could not verify endpoint - check URL and deployment name'})
    
    except Exception as e:
        logger.error(f"Azure OpenAI verification failed: {e}")
        return jsonify({'valid': False, 'error': str(e)})

@app.route('/api/settings/verify/azure-storage', methods=['POST'])
def verify_azure_storage():
    """Verify Azure Storage connection string"""
    try:
        data = request.get_json()
        connection_string = data.get('connectionString', '').strip()
        
        if not connection_string:
            return jsonify({'valid': False, 'error': 'Connection string is required'})
        
        try:
            from azure.storage.blob import BlobServiceClient
            
            blob_service_client = BlobServiceClient.from_connection_string(connection_string)
            
            # Get account info to verify connection (doesn't require listing containers)
            try:
                account_info = blob_service_client.get_account_information()
                return jsonify({
                    'valid': True,
                    'message': 'Connection successful',
                    'sku': account_info.get('sku_name', 'Unknown')
                })
            except Exception:
                # Fallback: try to list containers without max_results parameter
                container_count = 0
                for container in blob_service_client.list_containers():
                    container_count += 1
                    if container_count >= 5:  # Only count up to 5
                        break
                
                return jsonify({
                    'valid': True,
                    'message': 'Connection successful',
                    'containers': container_count
                })
        
        except ImportError:
            return jsonify({
                'valid': False,
                'error': 'Azure Storage SDK not installed. Run: pip install azure-storage-blob'
            })
        except Exception as e:
            error_msg = str(e)
            if 'AuthenticationFailed' in error_msg:
                return jsonify({'valid': False, 'error': 'Authentication failed - check connection string'})
            elif 'getaddrinfo failed' in error_msg:
                return jsonify({'valid': False, 'error': 'Could not resolve storage account name'})
            elif 'InvalidResourceName' in error_msg:
                return jsonify({'valid': False, 'error': 'Invalid storage account name in connection string'})
            else:
                return jsonify({'valid': False, 'error': f'Connection failed: {error_msg}'})
    
    except Exception as e:
        logger.error(f"Azure Storage verification failed: {e}")
        return jsonify({'valid': False, 'error': str(e)})

@app.route('/api/settings/verify/path', methods=['POST'])
def verify_path():
    """Verify paths exist or create them if requested"""
    try:
        data = request.get_json()
        paths = data.get('paths', [])
        create_if_missing = data.get('createIfMissing', False)
        
        if not paths:
            return jsonify({'valid': False, 'error': 'No paths provided'})
        
        created = False
        for path_info in paths:
            path_value = path_info.get('value', '').strip()
            
            if not path_value:
                continue
            
            # Expand environment variables and user paths
            expanded_path = os.path.expandvars(os.path.expanduser(path_value))
            
            if not os.path.exists(expanded_path):
                if create_if_missing:
                    try:
                        os.makedirs(expanded_path, exist_ok=True)
                        created = True
                        logger.info(f"Created directory: {expanded_path}")
                    except Exception as e:
                        return jsonify({
                            'valid': False,
                            'error': f"Failed to create {path_value}: {str(e)}"
                        })
                else:
                    return jsonify({
                        'valid': False,
                        'error': f"Path does not exist: {path_value}"
                    })
            elif not os.path.isdir(expanded_path):
                return jsonify({
                    'valid': False,
                    'error': f"Path is not a directory: {path_value}"
                })
        
        return jsonify({
            'valid': True,
            'message': 'All paths verified',
            'created': created
        })
    
    except Exception as e:
        logger.error(f"Path verification failed: {e}")
        return jsonify({'valid': False, 'error': str(e)})

# ============================================================================
# AZURE STORAGE API ENDPOINTS
# ============================================================================
@app.route('/api/azure/containers', methods=['GET'])
def list_azure_containers():
    """List all containers in Azure Blob Storage"""
    try:
        from azure.storage.blob import BlobServiceClient
        
        # Get connection string from environment
        connection_string = os.getenv('AZURE_STORAGE_CONNECTION_STRING', '')
        
        if not connection_string:
            return jsonify({
                'success': False,
                'error': 'Azure Storage not configured. Please set up in Settings.'
            })
        
        blob_service_client = BlobServiceClient.from_connection_string(connection_string)
        
        containers = []
        for container in blob_service_client.list_containers():
            containers.append({
                'name': container['name'],
                'last_modified': container.get('last_modified', '').isoformat() if container.get('last_modified') else None
            })
        
        return jsonify({
            'success': True,
            'containers': containers
        })
        
    except ImportError:
        return jsonify({
            'success': False,
            'error': 'Azure Storage SDK not installed. Run: pip install azure-storage-blob'
        })
    except Exception as e:
        logger.error(f"Failed to list containers: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        })

@app.route('/api/azure/blobs/<container_name>', methods=['GET'])
def list_azure_blobs(container_name):
    """List all blobs in a container"""
    try:
        from azure.storage.blob import BlobServiceClient
        
        connection_string = os.getenv('AZURE_STORAGE_CONNECTION_STRING', '')
        
        if not connection_string:
            return jsonify({
                'success': False,
                'error': 'Azure Storage not configured'
            })
        
        blob_service_client = BlobServiceClient.from_connection_string(connection_string)
        container_client = blob_service_client.get_container_client(container_name)
        
        blobs = []
        for blob in container_client.list_blobs():
            blobs.append({
                'name': blob.name,
                'size': blob.size,
                'last_modified': blob.last_modified.isoformat() if blob.last_modified else None,
                'content_type': blob.content_settings.content_type if blob.content_settings else None
            })
        
        return jsonify({
            'success': True,
            'container': container_name,
            'blobs': blobs
        })
        
    except Exception as e:
        logger.error(f"Failed to list blobs in {container_name}: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        })

@app.route('/api/azure/download/<container_name>/<path:blob_name>', methods=['GET'])
def download_azure_blob(container_name, blob_name):
    """Download a blob from Azure Storage"""
    try:
        from azure.storage.blob import BlobServiceClient
        from flask import Response
        import io
        
        connection_string = os.getenv('AZURE_STORAGE_CONNECTION_STRING', '')
        
        if not connection_string:
            return jsonify({
                'success': False,
                'error': 'Azure Storage not configured'
            }), 400
        
        blob_service_client = BlobServiceClient.from_connection_string(connection_string)
        blob_client = blob_service_client.get_blob_client(container_name, blob_name)
        
        # Download blob content
        blob_data = blob_client.download_blob()
        content = blob_data.readall()
        
        # Determine content type
        content_type_map = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.csv': 'text/csv',
            '.eml': 'message/rfc822',
            '.png': 'image/png',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg'
        }
        
        ext = os.path.splitext(blob_name)[1].lower()
        content_type = content_type_map.get(ext, 'application/octet-stream')
        
        return Response(
            content,
            mimetype=content_type,
            headers={
                'Content-Disposition': f'attachment; filename="{os.path.basename(blob_name)}"'
            }
        )
        
    except Exception as e:
        logger.error(f"Failed to download blob {blob_name}: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# ============================================================================
# FILE PREVIEW API ENDPOINTS
# ============================================================================

# Store parsed EML data temporarily for attachment access
_parsed_eml_cache = {}

@app.route('/api/preview/parse', methods=['POST'])
def parse_file_for_preview():
    """Parse uploaded file and return preview-ready content"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file provided'})
        
        file = request.files['file']
        filename = file.filename
        ext = os.path.splitext(filename)[1].lower()
        
        if ext == '.eml':
            return parse_eml_file(file, filename)
        elif ext == '.docx':
            return parse_docx_file(file, filename)
        elif ext in ['.xlsx', '.xls']:
            return parse_excel_file(file, filename)
        elif ext == '.csv':
            return parse_csv_file(file, filename)
        else:
            return jsonify({
                'success': True,
                'type': 'unsupported',
                'filename': filename,
                'message': f'{ext.upper()} files use native preview'
            })
    
    except Exception as e:
        logger.error(f"File preview parsing failed: {e}")
        return jsonify({'success': False, 'error': str(e)})

def parse_eml_file(file, filename):
    """Parse EML file and extract email content + attachments"""
    import email
    from email import policy
    from email.utils import parsedate_to_datetime
    import base64
    import uuid
    
    try:
        content = file.read()
        msg = email.message_from_bytes(content, policy=policy.default)
        
        # Extract headers
        headers = {
            'from': str(msg.get('From', '')),
            'to': str(msg.get('To', '')),
            'cc': str(msg.get('Cc', '')),
            'subject': str(msg.get('Subject', '')),
            'date': str(msg.get('Date', ''))
        }
        
        # Extract body
        body_html = None
        body_text = None
        attachments = []
        
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                content_disposition = str(part.get('Content-Disposition', ''))
                
                # Check if it's an attachment
                if 'attachment' in content_disposition or part.get_filename():
                    att_filename = part.get_filename() or f'attachment_{len(attachments)}'
                    att_data = part.get_payload(decode=True)
                    if att_data:
                        attachments.append({
                            'index': len(attachments),
                            'filename': att_filename,
                            'content_type': content_type,
                            'size': len(att_data),
                            'data': base64.b64encode(att_data).decode('utf-8')
                        })
                elif content_type == 'text/html' and not body_html:
                    payload = part.get_payload(decode=True)
                    if payload:
                        body_html = payload.decode('utf-8', errors='replace')
                elif content_type == 'text/plain' and not body_text:
                    payload = part.get_payload(decode=True)
                    if payload:
                        body_text = payload.decode('utf-8', errors='replace')
        else:
            # Simple message
            content_type = msg.get_content_type()
            payload = msg.get_payload(decode=True)
            if payload:
                if content_type == 'text/html':
                    body_html = payload.decode('utf-8', errors='replace')
                else:
                    body_text = payload.decode('utf-8', errors='replace')
        
        # Generate cache key for attachment access
        cache_key = str(uuid.uuid4())
        _parsed_eml_cache[cache_key] = {
            'attachments': attachments,
            'timestamp': time.time()
        }
        
        # Clean old cache entries (older than 30 minutes)
        current_time = time.time()
        keys_to_delete = [k for k, v in _parsed_eml_cache.items() 
                         if current_time - v['timestamp'] > 1800]
        for k in keys_to_delete:
            del _parsed_eml_cache[k]
        
        return jsonify({
            'success': True,
            'type': 'eml',
            'filename': filename,
            'cacheKey': cache_key,
            'headers': headers,
            'bodyHtml': body_html,
            'bodyText': body_text,
            'attachments': [{
                'index': a['index'],
                'filename': a['filename'],
                'contentType': a['content_type'],
                'size': a['size']
            } for a in attachments]
        })
        
    except Exception as e:
        logger.error(f"EML parsing failed: {e}")
        return jsonify({'success': False, 'error': f'EML parsing failed: {str(e)}'})

def parse_docx_file(file, filename):
    """Parse DOCX file and convert to HTML"""
    try:
        from docx import Document
        from docx.shared import Inches
        import io
        
        content = file.read()
        doc = Document(io.BytesIO(content))
        
        html_parts = ['<div class="docx-content">']
        
        for para in doc.paragraphs:
            style = para.style.name if para.style else 'Normal'
            text = para.text.strip()
            
            if not text:
                html_parts.append('<p>&nbsp;</p>')
                continue
            
            # Map styles to HTML
            if 'Heading 1' in style:
                html_parts.append(f'<h1>{text}</h1>')
            elif 'Heading 2' in style:
                html_parts.append(f'<h2>{text}</h2>')
            elif 'Heading 3' in style:
                html_parts.append(f'<h3>{text}</h3>')
            elif 'Title' in style:
                html_parts.append(f'<h1 class="title">{text}</h1>')
            else:
                html_parts.append(f'<p>{text}</p>')
        
        # Handle tables
        for table in doc.tables:
            html_parts.append('<table class="docx-table">')
            for row in table.rows:
                html_parts.append('<tr>')
                for cell in row.cells:
                    html_parts.append(f'<td>{cell.text}</td>')
                html_parts.append('</tr>')
            html_parts.append('</table>')
        
        html_parts.append('</div>')
        
        return jsonify({
            'success': True,
            'type': 'docx',
            'filename': filename,
            'html': '\n'.join(html_parts)
        })
        
    except ImportError:
        return jsonify({
            'success': False,
            'error': 'python-docx not installed. Run: pip install python-docx'
        })
    except Exception as e:
        logger.error(f"DOCX parsing failed: {e}")
        return jsonify({'success': False, 'error': f'DOCX parsing failed: {str(e)}'})

def parse_excel_file(file, filename):
    """Parse Excel file and convert to HTML table"""
    try:
        import pandas as pd
        import io
        
        content = file.read()
        
        # Read Excel file
        df = pd.read_excel(io.BytesIO(content), sheet_name=None)
        
        html_parts = ['<div class="excel-content">']
        
        for sheet_name, sheet_df in df.items():
            html_parts.append(f'<h3 class="sheet-name">📊 {sheet_name}</h3>')
            # Convert to HTML with styling
            table_html = sheet_df.to_html(
                classes='excel-table',
                index=False,
                na_rep='',
                max_rows=100
            )
            html_parts.append(table_html)
        
        html_parts.append('</div>')
        
        return jsonify({
            'success': True,
            'type': 'xlsx',
            'filename': filename,
            'html': '\n'.join(html_parts)
        })
        
    except ImportError:
        return jsonify({
            'success': False,
            'error': 'pandas/openpyxl not installed'
        })
    except Exception as e:
        logger.error(f"Excel parsing failed: {e}")
        return jsonify({'success': False, 'error': f'Excel parsing failed: {str(e)}'})

def parse_csv_file(file, filename):
    """Parse CSV file and convert to HTML table"""
    try:
        import pandas as pd
        import io
        
        content = file.read().decode('utf-8', errors='replace')
        df = pd.read_csv(io.StringIO(content))
        
        table_html = df.to_html(
            classes='csv-table',
            index=False,
            na_rep='',
            max_rows=100
        )
        
        return jsonify({
            'success': True,
            'type': 'csv',
            'filename': filename,
            'html': f'<div class="csv-content">{table_html}</div>'
        })
        
    except Exception as e:
        logger.error(f"CSV parsing failed: {e}")
        return jsonify({'success': False, 'error': f'CSV parsing failed: {str(e)}'})

@app.route('/api/preview/attachment/<cache_key>/<int:index>', methods=['GET'])
def get_eml_attachment(cache_key, index):
    """Get a specific attachment from a parsed EML file"""
    try:
        from flask import Response
        import base64
        
        if cache_key not in _parsed_eml_cache:
            return jsonify({'success': False, 'error': 'Cache expired, please reload email'}), 404
        
        cache_entry = _parsed_eml_cache[cache_key]
        attachments = cache_entry['attachments']
        
        if index < 0 or index >= len(attachments):
            return jsonify({'success': False, 'error': 'Attachment not found'}), 404
        
        attachment = attachments[index]
        content = base64.b64decode(attachment['data'])
        
        return Response(
            content,
            mimetype=attachment['content_type'],
            headers={
                'Content-Disposition': f'inline; filename="{attachment["filename"]}"'
            }
        )
        
    except Exception as e:
        logger.error(f"Failed to get attachment: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/pdf/split', methods=['POST'])
def split_pdf_files():
    """
    Intelligently process PDF files:
    - Split PDFs with multiple invoices
    - Process single invoice/non-invoice PDFs normally
    Supports file upload and automatic data extraction
    """
    start_time = time.time()
    
    if pdf_splitter is None:
        # Check if it's a configuration issue or missing dependencies
        try:
            from pdf import BarcodeVendorPDFSplitter
            temp_processor = BarcodeVendorPDFSplitter()
            deps = temp_processor.check_dependencies()
            
            if deps['missing']:
                return jsonify({
                    'success': False,
                    'error': 'PDF processing libraries not installed',
                    'missing_libraries': deps['missing'],
                    'install_command': f"pip install {' '.join(deps['missing'])}",
                    'details': 'Please install the required libraries and restart the server'
                }), 500
            else:
                return jsonify({
                    'success': False,
                    'error': 'PDF processor initialization failed'
                }), 500
        except ImportError:
            return jsonify({
                'success': False,
                'error': 'PDF processor module not available. Check pdf.py file'
            }), 500
    
    try:
        file_paths = []
        temp_files = []
        temp_dir = None
        
        # Handle uploaded files
        if 'files' in request.files:
            uploaded_files = request.files.getlist('files')
            if not temp_dir:
                temp_dir = tempfile.mkdtemp()
            
            for file in uploaded_files:
                if file.filename and file.filename.lower().endswith('.pdf'):
                    safe_filename = os.path.basename(file.filename)
                    file_path = os.path.join(temp_dir, safe_filename)
                    file.save(file_path)
                    
                    if os.path.exists(file_path):
                        file_paths.append(file_path)
                        temp_files.append(file_path)
        
        # Handle JSON data with paths
        data = None
        if request.is_json:
            data = request.json
        elif request.form.get('data'):
            try:
                data = json.loads(request.form.get('data'))
            except:
                pass
        
        if data and 'filePaths' in data:
            paths = data['filePaths']
            if isinstance(paths, str):
                paths = [paths]
            
            for path in paths:
                path = path.strip()
                if os.path.exists(path) and path.lower().endswith('.pdf'):
                    file_paths.append(path)
        
        if not file_paths:
            return jsonify({'success': False, 'error': 'No valid PDF files found'}), 400
        
        print(f"\n{'='*80}")
        print(f"🔍 INTELLIGENT PDF PROCESSING: {len(file_paths)} PDF file(s)")
        print(f"   Output: {Config.SPLIT_PDF_OUTPUT_PATH}")
        print(f"   Method: {Config.SPLIT_PDF_METHOD}")
        print(f"{'='*80}\n")
        
        # Process each PDF file
        all_results = {}
        total_splits = 0
        single_documents = 0
        multi_invoice_documents = 0
        
        for pdf_path in file_paths:
            pdf_name = os.path.basename(pdf_path)
            logger.info(f"Processing PDF: {pdf_name}")
            
            # Read PDF content
            with open(pdf_path, 'rb') as f:
                pdf_content = f.read()
            
            # Use high-performance parallel processor
            result = pdf_splitter.process_and_extract_parallel(pdf_content, pdf_name)
            
            all_results[pdf_name] = result
            
            if result.get('error'):
                logger.error(f"Error processing {pdf_name}: {result.get('error')}")
                continue
            
            # Update stats
            if result.get('processing_mode') == 'direct':
                single_documents += 1
            elif result.get('processing_mode') == 'split':
                multi_invoice_documents += 1
                total_splits += len(result.get('split_files', []))
        
        processing_time = time.time() - start_time
        
        # Clean up temp files
        for temp_file in temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
            except:
                pass
        
        if temp_dir and os.path.exists(temp_dir):
            try:
                os.rmdir(temp_dir)
            except:
                pass
        
        print(f"\n{'='*80}")
        print(f"✅ INTELLIGENT PDF PROCESSING COMPLETE!")
        print(f"   Total PDFs: {len(file_paths)}")
        print(f"   Single Documents: {single_documents}")
        print(f"   Multi-Invoice Documents: {multi_invoice_documents}")
        print(f"   Total Split Files: {total_splits}")
        print(f"   Time: {processing_time:.2f}s")
        print(f"{'='*80}\n")
        
        return jsonify({
            'success': True,
            'results': all_results,
            'stats': {
                'totalPdfs': len(file_paths),
                'singleDocuments': single_documents,
                'multiInvoiceDocuments': multi_invoice_documents,
                'totalSplits': total_splits,
                'processingTime': round(processing_time, 2),
                'outputDirectory': Config.SPLIT_PDF_OUTPUT_PATH,
                'autoExtraction': Config.SPLIT_PDF_AUTO_EXTRACT
            },
            'timingReport': {
                'totalTime': f"{processing_time:.3f}",
                'startTime': datetime.fromtimestamp(start_time).strftime('%H:%M:%S'),
                'endTime': datetime.now().strftime('%H:%M:%S')
            }
        })
    
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"\n❌ Error in PDF processing: {error_details}\n")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/process/parallel', methods=['POST'])
def process_files_parallel():
    """
    NEW: High-performance parallel document processing endpoint.
    Uses ThreadPoolExecutor for concurrent API calls.
    Supports: PDF, DOCX, EML (with attachments), Excel, CSV
    """
    start_time = time.time()
    
    if parallel_processor is None:
        return jsonify({
            'success': False,
            'error': 'Parallel processor module not loaded. Install dependencies and restart server.'
        }), 500
    
    try:
        file_paths = []
        temp_files = []
        temp_dir = None
        
        # Handle uploaded files
        if 'files' in request.files:
            uploaded_files = request.files.getlist('files')
            if not temp_dir:
                temp_dir = tempfile.mkdtemp()
            
            for file in uploaded_files:
                if file.filename:
                    safe_filename = os.path.basename(file.filename)
                    file_path = os.path.join(temp_dir, safe_filename)
                    file.save(file_path)
                    
                    if os.path.exists(file_path):
                        file_paths.append(file_path)
                        temp_files.append(file_path)
        
        # Determine JSON data source
        data = None
        if request.is_json:
            data = request.json
        elif request.form.get('data'):
            try:
                data = json.loads(request.form.get('data'))
            except:
                pass
        
        # Process JSON data (Blobs, Local Paths) if available
        if data:
            # Handle blob storage files (download first)
            if 'blobFiles' in data and extractor:
                env_vars = extractor.load_and_validate_env()
                if env_vars and env_vars.get('AZURE_STORAGE_CONNECTION_STRING'):
                    blob_files = data['blobFiles']
                    downloaded_files = extractor.download_blobs_to_temp_files(
                        env_vars['AZURE_STORAGE_CONNECTION_STRING'],
                        blob_files
                    )
                    for file_info in downloaded_files:
                        file_paths.append(file_info["temp_path"])
                        temp_files.append(file_info["temp_path"])
            
            # Handle local file paths
            if 'filePaths' in data:
                paths = data['filePaths']
                if isinstance(paths, str):
                    paths = [paths]
                
                for path in paths:
                    path = path.strip()
                    if os.path.exists(path):
                        if os.path.isdir(path):
                            # Scan folder for files
                            folder_files = parallel_processor.ParallelProcessingEngine().scan_path(path)
                            file_paths.extend(folder_files)
                        else:
                            file_paths.append(path)
            
            # Handle folder path
            if 'folderPath' in data:
                folder_path = data['folderPath'].strip()
                if os.path.isdir(folder_path):
                    folder_files = parallel_processor.ParallelProcessingEngine().scan_path(folder_path)
                    file_paths.extend(folder_files)
        
        if not file_paths:
            return jsonify({'success': False, 'error': 'No valid files found'}), 400
        
        # SEPARATE PDF AND OTHER FILES
        pdf_files = []
        other_files = []
        
        for file_path in file_paths:
            if file_path.lower().endswith('.pdf'):
                pdf_files.append(file_path)
            else:
                other_files.append(file_path)
        
        print(f"\n{'='*80}")
        print(f"🚀 PARALLEL PROCESSING: {len(file_paths)} file(s)")
        print(f"   PDF files: {len(pdf_files)}")
        print(f"   Other files: {len(other_files)}")
        print(f"{'='*80}\n")
        
        # Process PDF files with BEST PDF processor
        pdf_results = None
        if pdf_files and pdf_processor:
            print(f"\n📄 Processing {len(pdf_files)} PDF file(s) with BEST PDF Processor...")
            
            pdf_start_time = time.time()
            all_pdf_results = []
            
            for pdf_path in pdf_files:
                try:
                    print(f"\n📄 Processing: {Path(pdf_path).name}")
                    
                    with open(pdf_path, 'rb') as f:
                        pdf_content = f.read()
                    
                    # Process with BEST splitting
                    result = pdf_processor_instance.process_and_extract_parallel(pdf_content, Path(pdf_path).name)
                    all_pdf_results.append(result)
                    
                    print(f"✓ Completed: {Path(pdf_path).name}")
                    
                except Exception as pdf_error:
                    print(f"❌ Error processing {Path(pdf_path).name}: {str(pdf_error)}")
                    all_pdf_results.append({
                        'error': str(pdf_error),
                        'original_filename': Path(pdf_path).name
                    })
            
            pdf_duration = time.time() - pdf_start_time
            successful_pdfs = [r for r in all_pdf_results if 'error' not in r]
            
            pdf_results = {
                'total_pdfs': len(pdf_files),
                'successful': len(successful_pdfs),
                'failed': len(all_pdf_results) - len(successful_pdfs),
                'processing_duration': pdf_duration,
                'results': all_pdf_results
            }
            
            print(f"\n✅ PDF Processing Complete!")
            print(f"   Successful: {len(successful_pdfs)}/{len(pdf_files)}")
            print(f"   Processing time: {pdf_duration:.2f}s")
        
        # Process other files with parallel processor
        other_results = None
        if other_files and parallel_processor:
            print(f"\n📄 Processing {len(other_files)} other file(s) with parallel processor...")
            
            config = parallel_processor.ProcessingConfig(
                max_threads=Config.MAX_THREADS,
                max_processes=Config.MAX_PROCESSES,
                retry_attempts=Config.RETRY_ATTEMPTS,
                retry_delay_base=Config.RETRY_DELAY_BASE
            )
            engine = parallel_processor.ParallelProcessingEngine(config)
            
            other_results = engine.process_files_parallel(other_files)
            
            print(f"✅ Other files processing complete!")
        
        # Combine results
        processing_time = time.time() - start_time
        
        # Build combined results from PDF and other processing
        results = {}
        metadata = {"total_documents_processed": 0, "total_fields_extracted": 0, "api_calls": {"total": 0}}
        
        # Add PDF results to combined results
        if pdf_results:
            for pdf_result in pdf_results.get('results', []):
                if 'error' not in pdf_result:
                    filename = pdf_result.get('original_filename', 'unknown.pdf')
                    results[filename] = pdf_result
                    metadata["total_documents_processed"] += len(pdf_result.get('split_files', []))
                    metadata["api_calls"]["total"] += pdf_result.get('api_calls', {}).get('document_intelligence', 0)
                    metadata["api_calls"]["total"] += pdf_result.get('api_calls', {}).get('openai', 0)
        
        # Add other results to combined results
        if other_results:
            for file_name, file_result in other_results.items():
                if file_name != '_EXTRACTION_METADATA':
                    results[file_name] = file_result
            
            # Add metadata from other processing
            other_metadata = other_results.get("_EXTRACTION_METADATA", {})
            metadata["total_documents_processed"] += other_metadata.get('total_documents_processed', 0)
            metadata["total_fields_extracted"] += other_metadata.get('total_fields_extracted', 0)
            metadata["api_calls"]["total"] += other_metadata.get('api_calls', {}).get('total', 0)
        
        # Add metadata to results
        results["_EXTRACTION_METADATA"] = metadata
        
        # Build detected files info
        detected_files = []
        for idx, file_path in enumerate(file_paths, 1):
            path_obj = Path(file_path)
            file_type = get_file_type(path_obj.name)  # Use the existing get_file_type function
            try:
                size_bytes = path_obj.stat().st_size
                if size_bytes < 1024:
                    size_str = f"{size_bytes} bytes"
                elif size_bytes < 1024 * 1024:
                    size_str = f"{size_bytes / 1024:.1f} KB"
                else:
                    size_str = f"{size_bytes / (1024 * 1024):.2f} MB"
            except:
                size_str = "Unknown"
                size_bytes = 0
            
            detected_files.append({
                "number": idx,
                "filename": path_obj.name,
                "type": file_type or "Unknown",
                "size": size_str,
                "size_bytes": size_bytes,
                "location": str(path_obj.parent)
            })
        
        print(f"\n{'='*80}")
        print(f"✅ PARALLEL PROCESSING COMPLETE!")
        print(f"   Documents: {metadata.get('total_documents_processed', 0)}")
        print(f"   Fields: {metadata.get('total_fields_extracted', 0)}")
        print(f"   Time: {processing_time:.2f}s")
        print(f"   API Calls: {metadata.get('api_calls', {}).get('total', 0)}")
        print(f"{'='*80}\n")
        
        # Save combined JSON output
        output_file = None
        if results:
            try:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                output_filename = f"parallel_processing_results_{timestamp}.json"
                output_file = os.path.join(JSON_FILES_FOLDER, output_filename)
                
                with open(output_file, 'w', encoding='utf-8') as f:
                    json.dump(results, f, indent=2, ensure_ascii=False, default=str)
                
                print(f"[JSON] Saved combined results: {output_filename}")
            except Exception as e:
                print(f"[ERROR] Failed to save combined JSON: {e}")
                output_file = None
        # === AUTO-LOAD TO REVIEW SESSION ===
        # Populate review_session so /review page auto-loads these files
        global review_session
        review_docs = []
        review_session['json_files'] = {}  # Reset json files mapping
        
        for file_path in file_paths:
            file_name = os.path.basename(file_path)
            file_type = get_file_type(file_name)
            file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
            file_location = os.path.dirname(file_path)
            
            # Get per-file metadata from results
            file_result = results.get(file_name, {})
            file_metadata = file_result.get('_metadata', {})
            
            review_docs.append({
                'path': file_path,
                'name': file_name,
                'type': file_type,
                'size': file_size,
                'location': file_location,
                'temp': file_path in temp_files,
                'is_pdf': file_type == 'PDF',
                'processing_time': file_metadata.get('processing_time', 0),
                'api_calls': file_metadata.get('api_calls', {}),
            })
            
            # Save individual JSON file for this document
            if file_name in results and file_name != '_EXTRACTION_METADATA':
                file_data = results[file_name]
                
                # Create individual JSON file
                json_filename = f"{Path(file_name).stem}_{int(time.time())}.json"
                json_path = os.path.join(JSON_FILES_FOLDER, json_filename)
                
                try:
                    with open(json_path, 'w', encoding='utf-8') as f:
                        json.dump({
                            'source_file': file_name,
                            'source_path': file_path,
                            'file_type': file_type,
                            'extraction_timestamp': datetime.now().isoformat(),
                            'data': file_data
                        }, f, indent=2, ensure_ascii=False)
                    
                    review_session['json_files'][file_name] = json_path
                    print(f"[JSON] Saved: {json_filename}")
                except Exception as e:
                    print(f"[ERROR] Failed to save JSON for {file_name}: {e}")
                
                # Also load extracted fields for review
                fields = {}
                for category, category_data in file_data.items():
                    if category == '_metadata':
                        continue
                    if isinstance(category_data, dict):
                        for field_name, field_info in category_data.items():
                            if isinstance(field_info, dict) and 'value' in field_info:
                                fields[field_name] = {
                                    'value': field_info.get('value', ''),
                                    'type': field_info.get('type', 'string'),
                                    'confidence': field_info.get('confidence', 0.0),
                                    'page': field_info.get('page', 'N/A'),
                                    'source': category
                                }
                review_session['extracted_data'][file_name] = fields
        
        review_session['documents'] = review_docs
        review_session['current_index'] = 0
        print(f"[AUTO-REVIEW] Loaded {len(review_docs)} document(s) into review session")
        print(f"[JSON FILES] Saved {len(review_session['json_files'])} JSON files to {JSON_FILES_FOLDER}")

        # Clean up temp files
        # Only delete files that are NOT in the review session
        current_review_paths = [d['path'] for d in review_docs]
        
        for temp_file in temp_files:
            try:
                if temp_file not in current_review_paths:
                    if os.path.exists(temp_file):
                        os.remove(temp_file)
            except:
                pass
        
        return jsonify({
            'success': True,
            'detectedFiles': detected_files,
            'data': results,
            'outputFile': output_file,
            'stats': {
                'totalFiles': len(file_paths),
                'pdfFiles': len(pdf_files),
                'otherFiles': len(other_files),
                'totalFieldsExtracted': metadata.get('total_fields_extracted', 0),
                'successfulExtractions': metadata.get('total_documents_processed', 0),
                'failedExtractions': 0,  # Calculate if needed
                'processingTime': round(processing_time, 2),
                'documentIntelligenceCalls': metadata.get('api_calls', {}).get('document_intelligence', 0),
                'openaiCalls': metadata.get('api_calls', {}).get('openai', 0),
                'totalApiCalls': metadata.get('api_calls', {}).get('total', 0),
                'detailedFiles': review_docs  # Per-file details for frontend table
            },
            'timingReport': {
                'totalTime': f"{processing_time:.3f}",
                'startTime': datetime.fromtimestamp(start_time).strftime('%H:%M:%S'),
                'endTime': datetime.now().strftime('%H:%M:%S')
            },
            'metadata': {
                'extraction_mode': 'Parallel Processing with BEST PDF Splitting',
                'max_threads': Config.MAX_THREADS,
                'max_processes': Config.MAX_PROCESSES,
                'retry_attempts': Config.RETRY_ATTEMPTS,
                'pdf_processor': 'BarcodeVendorPDFSplitter' if pdf_processor else 'None'
            },
            'pdfResults': pdf_results,
            'otherResults': other_results
        })
    
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"\n❌ Error in parallel processing: {error_details}\n")
        return jsonify({'success': False, 'error': str(e)}), 500
    
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"\n❌ Error in parallel processing: {error_details}\n")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/blob/containers', methods=['GET'])
def list_containers():
    """List all Azure Blob Storage containers."""
    if extractor is None:
        return jsonify({'success': False, 'error': 'Extractor not loaded'}), 500
    
    try:
        env_vars = extractor.load_and_validate_env()
        if not env_vars or not env_vars.get('AZURE_STORAGE_CONNECTION_STRING'):
            return jsonify({'success': False, 'error': 'Azure Blob Storage not configured'}), 400
        
        if not extractor.BLOB_STORAGE_AVAILABLE:
            return jsonify({'success': False, 'error': 'Azure Blob Storage library not available'}), 400
        
        containers = extractor.list_blob_containers(env_vars['AZURE_STORAGE_CONNECTION_STRING'])
        return jsonify({'success': True, 'containers': containers})
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/blob/browse', methods=['POST'])
def browse_blob_path():
    """Browse files and folders in Azure Blob Storage at a given path."""
    if extractor is None:
        return jsonify({'success': False, 'error': 'Extractor not loaded'}), 500
    
    try:
        data = request.json
        container = data.get('container')
        prefix = data.get('prefix', '')
        
        env_vars = extractor.load_and_validate_env()
        if not env_vars or not env_vars.get('AZURE_STORAGE_CONNECTION_STRING'):
            return jsonify({'success': False, 'error': 'Azure Blob Storage not configured'}), 400
        
        folders, files = extractor.list_blob_items_at_path(
            env_vars['AZURE_STORAGE_CONNECTION_STRING'],
            container,
            prefix
        )
        
        # Format file sizes
        for file in files:
            size_bytes = file['size']
            if size_bytes < 1024:
                file['size_display'] = f"{size_bytes} bytes"
            elif size_bytes < 1024 * 1024:
                file['size_display'] = f"{size_bytes / 1024:.1f} KB"
            else:
                file['size_display'] = f"{size_bytes / (1024 * 1024):.2f} MB"
        
        return jsonify({
            'success': True,
            'folders': folders,
            'files': files,
            'current_path': f"{container}/{prefix}" if prefix else container
        })
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/local/browse', methods=['POST'])
def browse_local_path():
    """Browse files and folders in local file system."""
    if extractor is None:
        return jsonify({'success': False, 'error': 'Extractor not loaded'}), 500
    
    try:
        data = request.json
        path = data.get('path')
        
        # If path is None or empty, use home directory
        if not path:
            path = os.path.expanduser('~')
        
        # Security: Validate path to prevent traversal attacks
        try:
            path_obj = validate_path(path)
        except ValidationError as e:
            logger.warning(f"Path validation failed: {path} - {e.message}")
            return jsonify({'success': False, 'error': e.message}), 400
        
        if not path_obj.exists():
            return jsonify({'success': False, 'error': 'Path does not exist'}), 400
        
        if not path_obj.is_dir():
            return jsonify({'success': False, 'error': 'Path is not a directory'}), 400
        
        folders = []
        files = []
        supported_extensions = {'.eml', '.pdf', '.docx', '.doc', '.xlsx', '.xls', '.csv'}
        
        try:
            for item in path_obj.iterdir():
                if item.is_dir():
                    folders.append({
                        'name': item.name,
                        'path': str(item),
                        'type': 'folder'
                    })
                elif item.is_file() and item.suffix.lower() in supported_extensions:
                    size_bytes = item.stat().st_size
                    if size_bytes < 1024:
                        size_display = f"{size_bytes} bytes"
                    elif size_bytes < 1024 * 1024:
                        size_display = f"{size_bytes / 1024:.1f} KB"
                    else:
                        size_display = f"{size_bytes / (1024 * 1024):.2f} MB"
                    
                    files.append({
                        'name': item.name,
                        'path': str(item),
                        'size': size_bytes,
                        'size_display': size_display,
                        'type': 'file'
                    })
        except PermissionError:
            return jsonify({'success': False, 'error': 'Access denied'}), 403
        
        # Sort folders and files
        folders.sort(key=lambda x: x['name'].lower())
        files.sort(key=lambda x: x['name'].lower())
        
        # Get parent directory
        parent = str(path_obj.parent) if path_obj.parent != path_obj else None
        
        return jsonify({
            'success': True,
            'folders': folders,
            'files': files,
            'current_path': str(path_obj),
            'parent_path': parent
        })
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/process', methods=['POST'])
def process_files():
    start_time = time.time()
    
    if extractor is None:
        return jsonify({
            'success': False,
            'error': 'Extractor module not loaded. Check server console.'
        }), 500
    
    try:
        file_paths = []
        temp_files = []
        temp_dir = None
        blob_files_info = []
        
        # Handle uploaded files
        if 'files' in request.files:
            uploaded_files = request.files.getlist('files')
            temp_dir = tempfile.mkdtemp()
            
            for file in uploaded_files:
                if file.filename:
                    safe_filename = os.path.basename(file.filename)
                    file_path = os.path.join(temp_dir, safe_filename)
                    file.save(file_path)
                    
                    if os.path.exists(file_path):
                        file_paths.append(file_path)
                        temp_files.append(file_path)
        
        # Handle JSON data with paths or blob info
        elif request.is_json:
            data = request.json
            
            # Handle blob storage files
            if 'blobFiles' in data:
                env_vars = extractor.load_and_validate_env()
                if not env_vars or not env_vars.get('AZURE_STORAGE_CONNECTION_STRING'):
                    return jsonify({'success': False, 'error': 'Azure Blob Storage not configured'}), 400
                
                blob_files = data['blobFiles']
                
                # Download blobs
                downloaded_files = extractor.download_blobs_to_temp_files(
                    env_vars['AZURE_STORAGE_CONNECTION_STRING'],
                    blob_files
                )
                
                for file_info in downloaded_files:
                    file_paths.append(file_info["temp_path"])
                    temp_files.append(file_info["temp_path"])
                    blob_files_info.append({
                        'filename': Path(file_info["original_name"]).name,
                        'location': f"Azure Blob: {file_info['blob_info']['container']}/{file_info['original_name']}",
                        'size': file_info['blob_info']['size']
                    })
            
            # Handle local file paths
            if 'filePaths' in data:
                paths = data['filePaths']
                if isinstance(paths, str):
                    paths = [paths]
                
                for path in paths:
                    path = path.strip()
                    if os.path.exists(path):
                        if os.path.isdir(path):
                            folder_files = extractor.scan_folder_for_files(path)
                            if folder_files:
                                file_paths.extend(folder_files)
                        else:
                            file_paths.append(path)
        else:
            return jsonify({'success': False, 'error': 'No files or paths provided'}), 400
        
        if not file_paths:
            return jsonify({'success': False, 'error': 'No valid files found'}), 400
        
        # Load environment variables
        env_vars = extractor.load_and_validate_env()
        if not env_vars:
            return jsonify({'success': False, 'error': 'Environment variables not configured'}), 500
        
        # SEPARATE PDF AND OTHER FILES
        pdf_files, other_files = extractor.separate_pdf_and_other_files(file_paths)
        
        # Reset global trackers
        extractor.api_call_tracker = {"document_intelligence": 0, "azure_openai": 0}
        extractor.timing_data = {
            "process_start_time": start_time,
            "process_end_time": None,
            "folder_scan_start": None,
            "folder_scan_end": None,
            "file_detection_complete": time.time(),
            "parsing_start_time": None,
            "parsing_complete_time": None,
            "doc_intelligence_start": None,
            "doc_intelligence_end": None,
            "ai_extraction_start": None,
            "ai_extraction_end": None,
            "json_save_time": None,
            "file_timings": {}
        }
        
        # PROCESS PDF FILES WITH pdf.py
        pdf_results = None
        if pdf_files and pdf_processor:
            print(f"\n{'='*80}")
            print(f"Processing {len(pdf_files)} PDF file(s) with BEST PDF Processor...")
            print(f"{'='*80}\n")
            
            pdf_start_time = time.time()
            
            try:
                # Use the already initialized processor
                processor = pdf_processor_instance
                
                # Process each PDF individually (simpler and more reliable)
                all_pdf_results = []
                for pdf_path in pdf_files:
                    try:
                        print(f"\n📄 Processing: {Path(pdf_path).name}")
                        
                        with open(pdf_path, 'rb') as f:
                            pdf_content = f.read()
                        
                        # Process with BEST splitting
                        result = processor.process_and_extract_parallel(pdf_content, Path(pdf_path).name)
                        all_pdf_results.append(result)
                        
                        print(f"✓ Completed: {Path(pdf_path).name}")
                        
                    except Exception as pdf_error:
                        print(f"❌ Error processing {Path(pdf_path).name}: {str(pdf_error)}")
                        all_pdf_results.append({
                            'error': str(pdf_error),
                            'original_filename': Path(pdf_path).name
                        })
                
                pdf_end_time = time.time()
                pdf_duration = pdf_end_time - pdf_start_time
                
                # Aggregate results
                successful = [r for r in all_pdf_results if 'error' not in r]
                failed = [r for r in all_pdf_results if 'error' in r]
                total_splits = sum(len(r.get('split_files', [])) for r in successful)
                total_api_calls = sum(r.get('api_calls', {}).get('document_intelligence', 0) for r in successful)
                total_openai_calls = sum(r.get('api_calls', {}).get('openai', 0) for r in successful)
                
                pdf_results = {
                    'total_pdfs': len(pdf_files),
                    'pdf_paths': [f"Path {idx}: {Path(p).name}" for idx, p in enumerate(pdf_files, 1)],
                    'successful': len(successful),
                    'failed': len(failed),
                    'total_output_files': total_splits,
                    'processing_duration': pdf_duration,
                    'split_pdfs_directory': processor.output_pdf_dir,
                    'json_outputs_directory': processor.output_json_dir,
                    'api_calls': {
                        'document_intelligence': total_api_calls,
                        'openai': total_openai_calls
                    },
                    'results': all_pdf_results
                }
                
                print(f"\n{'='*80}")
                print(f"✅ PDF Processing Complete!")
                print(f"   Successful: {len(successful)}/{len(pdf_files)}")
                print(f"   Total output files: {total_splits}")
                print(f"   Processing time: {pdf_duration:.2f}s")
                print(f"   Document Intelligence calls: {total_api_calls}")
                print(f"   OpenAI calls: {total_openai_calls}")
                print(f"{'='*80}\n")
                
            except Exception as pdf_error:
                import traceback
                error_trace = traceback.format_exc()
                print(f"\n❌ PDF Processing Error: {str(pdf_error)}")
                print(f"Traceback:\n{error_trace}\n")
                
                # Return error but don't fail the whole request
                pdf_results = {
                    'error': str(pdf_error),
                    'error_trace': error_trace,
                    'total_pdfs': len(pdf_files),
                    'pdf_paths': [f"Path {idx}: {Path(p).name}" for idx, p in enumerate(pdf_files, 1)],
                    'note': 'PDF processing failed, but other files may have been processed successfully'
                }
        
        # Collect detected files info (including PDFs)
        detected_files = []
        for idx, file_path in enumerate(file_paths, 1):
            path_obj = Path(file_path)
            file_type = extractor.detect_file_type(file_path)
            
            try:
                size_bytes = path_obj.stat().st_size
                if size_bytes < 1024:
                    size_str = f"{size_bytes} bytes"
                elif size_bytes < 1024 * 1024:
                    size_str = f"{size_bytes / 1024:.1f} KB"
                else:
                    size_str = f"{size_bytes / (1024 * 1024):.2f} MB"
            except:
                size_str = "Unknown"
                size_bytes = 0
            
            # Check if this is a blob file
            blob_info = next((b for b in blob_files_info if b['filename'] == path_obj.name), None)
            location = blob_info['location'] if blob_info else str(path_obj.parent)
            
            detected_files.append({
                "number": idx,
                "filename": path_obj.name,
                "type": file_type or "Unknown",
                "size": size_str,
                "size_bytes": size_bytes,
                "location": location,
                "is_pdf": file_type == 'PDF'
            })
        
        # PROCESS OTHER FILES WITH multi_file_extractor
        structured_data = None
        if other_files:
            print(f"\n{'='*80}")
            print(f"Processing {len(other_files)} non-PDF file(s) with multi_file_extractor...")
            print(f"{'='*80}\n")
            
            # Collect all documents (excluding PDFs)
            all_eml_data, all_documents = extractor.collect_all_documents(other_files)
            
            if not all_eml_data and not all_documents:
                print("⚠️  No non-PDF documents to process")
            else:
                # Analyze documents
                extracted_data = extractor.analyze_documents_single_call(
                    all_eml_data,
                    all_documents,
                    env_vars["DOC_ENDPOINT"],
                    env_vars["DOC_KEY"]
                )
                
                if not extracted_data:
                    return jsonify({'success': False, 'error': 'Failed to analyze documents'}), 500
                
                # Extract fields
                structured_data = extractor.extract_fields_with_file_sections(
                    extracted_data,
                    env_vars["AZURE_OPENAI_ENDPOINT"],
                    env_vars["AZURE_OPENAI_KEY"]
                )
                
                if not structured_data:
                    return jsonify({'success': False, 'error': 'Failed to extract fields'}), 500
                
                # Save JSON output
                output_dir = r"C:\Users\isarkar2\OneDrive - DXC Production\Desktop\PDF_Console_App\Output Json"
                output_file = extractor.save_json_output(structured_data, output_dir)
                
                if not output_file:
                    return jsonify({'success': False, 'error': 'Failed to save JSON output'}), 500
        
        # Record end time
        extractor.timing_data["process_end_time"] = time.time()
        total_time = extractor.timing_data["process_end_time"] - extractor.timing_data["process_start_time"]
        
        # Build timing report
        start_dt = datetime.fromtimestamp(extractor.timing_data["process_start_time"])
        end_dt = datetime.fromtimestamp(extractor.timing_data["process_end_time"])
        
        timing_stages = []
        stage_definitions = [
            ("File Detection", extractor.timing_data.get("process_start_time"), extractor.timing_data.get("file_detection_complete")),
            ("Document Parsing", extractor.timing_data.get("parsing_start_time"), extractor.timing_data.get("parsing_complete_time")),
            ("Document Intelligence", extractor.timing_data.get("doc_intelligence_start"), extractor.timing_data.get("doc_intelligence_end")),
            ("AI Field Extraction", extractor.timing_data.get("ai_extraction_start"), extractor.timing_data.get("ai_extraction_end")),
            ("JSON Saving", extractor.timing_data.get("json_save_time"), extractor.timing_data.get("process_end_time"))
        ]
        
        for stage_name, stage_start, stage_end in stage_definitions:
            if stage_start and stage_end:
                duration = stage_end - stage_start
                percentage = (duration / total_time) * 100
                timing_stages.append({
                    "name": stage_name,
                    "duration": f"{duration:.3f}",
                    "percentage": f"{percentage:.1f}"
                })
        
        timing_report = {
            "startTime": start_dt.strftime('%H:%M:%S.%f')[:-3],
            "endTime": end_dt.strftime('%H:%M:%S.%f')[:-3],
            "totalTime": f"{total_time:.3f}",
            "stages": timing_stages
        }
        
        # Extract metadata for tokens and API calls
        metadata = structured_data.get("_EXTRACTION_METADATA", {}) if structured_data else {}
        
        # Clean up temp files
        for temp_file in temp_files:
            try:
                os.remove(temp_file)
            except:
                pass
        
        if temp_dir and os.path.exists(temp_dir):
            try:
                os.rmdir(temp_dir)
            except:
                pass
        
        # Build response
        response_data = {
            'success': True,
            'detectedFiles': detected_files,
            'timingReport': timing_report,
            'stats': {
                'totalFiles': len(file_paths),
                'pdfFiles': len(pdf_files),
                'otherFiles': len(other_files),
                'totalApiCalls': extractor.api_call_tracker['document_intelligence'] + extractor.api_call_tracker['azure_openai'],
                'documentIntelligenceCalls': extractor.api_call_tracker['document_intelligence'],
                'azureOpenAICalls': extractor.api_call_tracker['azure_openai']
            }
        }
        
        # Add PDF results if available
        if pdf_results:
            response_data['pdfResults'] = pdf_results
        
        # Add other file results if available
        if structured_data:
            response_data['data'] = structured_data
            response_data['outputFile'] = output_file if 'output_file' in locals() else None
            response_data['metadata'] = {
                'extraction_mode': metadata.get('extraction_mode', 'Unknown'),
                'batch_size': metadata.get('batch_size', 0),
                'total_batches': metadata.get('total_batches', 0),
                'model': metadata.get('extraction_parameters', {}).get('model', 'Unknown'),
                'max_tokens': metadata.get('extraction_parameters', {}).get('max_tokens', 0)
            }
        
        return jsonify(response_data)
    
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"\n[ERROR] {error_details}\n")
        return jsonify({'success': False, 'error': str(e)}), 500


# ============================================================================
# DOCUMENT REVIEW API ENDPOINTS
# ============================================================================

@app.route('/api/review/config', methods=['GET'])
def review_get_config():
    """Return document review configuration paths"""
    return jsonify({
        'success': True,
        'acceptedFolder': ACCEPTED_FOLDER,
        'rejectedFolder': REJECTED_FOLDER,
        'processorAvailable': parallel_processor is not None
    })


@app.route('/api/review/documents', methods=['GET'])
def review_list_documents():
    """List all documents in current review session"""
    return jsonify({
        'success': True,
        'documents': review_session['documents'],
        'currentIndex': review_session['current_index'],
        'total': len(review_session['documents'])
    })


@app.route('/api/review/upload', methods=['POST'])
def review_upload_documents():
    """Upload documents for review"""
    global review_session
    
    try:
        file_paths = []
        
        # Handle file uploads
        if 'files' in request.files:
            uploaded_files = request.files.getlist('files')
            temp_dir = tempfile.mkdtemp()
            
            for file in uploaded_files:
                if file.filename:
                    safe_filename = os.path.basename(file.filename)
                    file_path = os.path.join(temp_dir, safe_filename)
                    file.save(file_path)
                    
                    if os.path.exists(file_path):
                        file_type = get_file_type(safe_filename)
                        file_paths.append({
                            'path': file_path,
                            'name': safe_filename,
                            'type': file_type,
                            'size': os.path.getsize(file_path),
                            'temp': True
                        })
        
        # Handle JSON paths
        elif request.is_json:
            data = request.json
            paths = data.get('paths', [])
            
            for path in paths:
                if os.path.isfile(path):
                    file_paths.append({
                        'path': path,
                        'name': os.path.basename(path),
                        'type': get_file_type(path),
                        'size': os.path.getsize(path),
                        'temp': False
                    })
                elif os.path.isdir(path):
                    for root, dirs, files in os.walk(path):
                        for f in files:
                            fp = os.path.join(root, f)
                            ft = get_file_type(f)
                            if ft != 'UNKNOWN':
                                file_paths.append({
                                    'path': fp,
                                    'name': f,
                                    'type': ft,
                                    'size': os.path.getsize(fp),
                                    'temp': False
                                })
        
        review_session['documents'] = file_paths
        review_session['extracted_data'] = {}
        review_session['current_index'] = 0
        
        return jsonify({
            'success': True,
            'count': len(file_paths),
            'documents': file_paths
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/review/add', methods=['POST'])
def review_add_documents():
    """Add more documents to current review session"""
    global review_session
    
    try:
        existing = review_session.get('documents', [])
        new_docs = []
        
        if 'files' in request.files:
            uploaded_files = request.files.getlist('files')
            temp_dir = tempfile.mkdtemp()
            
            for file in uploaded_files:
                if file.filename:
                    safe_filename = os.path.basename(file.filename)
                    file_path = os.path.join(temp_dir, safe_filename)
                    file.save(file_path)
                    
                    if os.path.exists(file_path):
                        new_docs.append({
                            'path': file_path,
                            'name': safe_filename,
                            'type': get_file_type(safe_filename),
                            'size': os.path.getsize(file_path),
                            'temp': True
                        })
        
        review_session['documents'] = existing + new_docs
        
        return jsonify({
            'success': True,
            'added': len(new_docs),
            'total': len(review_session['documents'])
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/review/document/<int:index>', methods=['GET'])
def review_get_document(index):
    """Get a specific document by index"""
    try:
        docs = review_session.get('documents', [])
        
        if index < 0 or index >= len(docs):
            return jsonify({'success': False, 'error': 'Invalid index'}), 400
        
        doc = docs[index]
        review_session['current_index'] = index
        
        return jsonify({
            'success': True,
            'document': doc,
            'index': index,
            'total': len(docs)
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/review/document/<int:index>/file', methods=['GET'])
def review_get_document_file(index):
    """Serve the actual document file"""
    try:
        docs = review_session.get('documents', [])
        
        if index < 0 or index >= len(docs):
            return jsonify({'success': False, 'error': 'Invalid index'}), 400
        
        doc = docs[index]
        return send_file(doc['path'], as_attachment=False)
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/review/document/<int:index>/content', methods=['GET'])
def review_get_document_content(index):
    """Get document content as text for preview"""
    try:
        docs = review_session.get('documents', [])
        
        if index < 0 or index >= len(docs):
            return jsonify({'success': False, 'error': 'Invalid index'}), 400
        
        doc = docs[index]
        file_type = doc['type']
        file_path = doc['path']
        
        content = ""
        
        if file_type == 'CSV':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
        elif file_type == 'EML':
            from email import policy
            from email.parser import BytesParser
            with open(file_path, 'rb') as f:
                msg = BytesParser(policy=policy.default).parse(f)
                content = f"From: {msg.get('From', '')}\n"
                content += f"To: {msg.get('To', '')}\n"
                content += f"Subject: {msg.get('Subject', '')}\n"
                content += f"Date: {msg.get('Date', '')}\n\n"
                
                if msg.is_multipart():
                    for part in msg.walk():
                        if part.get_content_type() == 'text/plain':
                            try:
                                content += part.get_content()
                            except:
                                pass
                else:
                    try:
                        content += msg.get_content()
                    except:
                        pass
                
                attachments = []
                if msg.is_multipart():
                    for part in msg.walk():
                        if 'attachment' in str(part.get('Content-Disposition', '')):
                            att_name = part.get_filename()
                            if att_name:
                                attachments.append(att_name)
                
                if attachments:
                    content += f"\n\n--- Attachments ({len(attachments)}) ---\n"
                    for att in attachments:
                        content += f"  - {att}\n"
        
        return jsonify({
            'success': True,
            'content': content,
            'type': file_type
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/review/document/<int:index>/extract', methods=['POST'])
def review_extract_fields(index):
    """Extract fields from a document using parallel processor"""
    try:
        docs = review_session.get('documents', [])
        
        if index < 0 or index >= len(docs):
            return jsonify({'success': False, 'error': 'Invalid index'}), 400
        
        doc = docs[index]
        doc_name = doc['name']
        
        # Check if already extracted
        if doc_name in review_session['extracted_data']:
            return jsonify({
                'success': True,
                'fields': review_session['extracted_data'][doc_name],
                'cached': True
            })
        
        if parallel_processor is None:
            return jsonify({
                'success': False, 
                'error': 'Parallel processor not available'
            }), 500
        
        # Extract using parallel processor (single document)
        config = parallel_processor.ProcessingConfig(
            max_threads=1,
            max_processes=1,
            retry_attempts=3
        )
        engine = parallel_processor.ParallelProcessingEngine(config)
        
        result = engine.process_single_document(doc['path'])
        
        fields = {}
        if result.success:
            for field_name, field_obj in result.fields.items():
                fields[field_name] = {
                    'value': field_obj.value,
                    'type': field_obj.field_type,
                    'confidence': field_obj.confidence,
                    'source': field_obj.source
                }
            
            for att_idx, att in enumerate(result.attachments, 1):
                for field_name, field_obj in att.fields.items():
                    key = f"Attachment_{att_idx}_{field_name}"
                    fields[key] = {
                        'value': field_obj.value,
                        'type': field_obj.field_type,
                        'confidence': field_obj.confidence,
                        'source': f"Attachment: {att.filename}"
                    }
        
        review_session['extracted_data'][doc_name] = fields
        
        return jsonify({
            'success': True,
            'fields': fields,
            'cached': False,
            'processingTime': result.processing_time
        })
        
    except Exception as e:
        import traceback
        print(f"Extraction error: {traceback.format_exc()}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/review/document/<int:index>/fields', methods=['PUT'])
def review_update_fields(index):
    """Update extracted fields for a document"""
    try:
        docs = review_session.get('documents', [])
        
        if index < 0 or index >= len(docs):
            return jsonify({'success': False, 'error': 'Invalid index'}), 400
        
        doc = docs[index]
        data = request.json
        fields = data.get('fields', {})
        
        review_session['extracted_data'][doc['name']] = fields
        
        return jsonify({
            'success': True,
            'message': 'Fields updated'
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/review/document/<int:index>/approve', methods=['POST'])
def review_approve_document(index):
    """Approve document and save to ACCEPTED folder"""
    try:
        docs = review_session.get('documents', [])
        
        if index < 0 or index >= len(docs):
            return jsonify({'success': False, 'error': 'Invalid index'}), 400
        
        doc = docs[index]
        data = request.json
        fields = data.get('fields', review_session['extracted_data'].get(doc['name'], {}))
        reviewer = data.get('reviewer', 'Anonymous')
        notes = data.get('notes', '')
        
        output_data = {
            'document_name': doc['name'],
            'document_type': doc['type'],
            'document_path': doc.get('path', ''),
            'document_location': doc.get('location', ''),
            'document_size': doc.get('size', 0),
            'review_status': 'APPROVED',
            'reviewer': reviewer,
            'review_notes': notes,
            'review_timestamp': datetime.now().isoformat(),
            'processing_time': doc.get('processing_time', 0),
            'api_calls': doc.get('api_calls', {}),
            'extracted_fields': fields
        }
        
        base_name = Path(doc['name']).stem
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_filename = f"{base_name}_APPROVED_{timestamp}.json"
        output_path = os.path.join(ACCEPTED_FOLDER, output_filename)
        
        counter = 1
        while os.path.exists(output_path):
            output_filename = f"{base_name}_APPROVED_{timestamp}_{counter}.json"
            output_path = os.path.join(ACCEPTED_FOLDER, output_filename)
            counter += 1
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(output_data, f, indent=2, ensure_ascii=False, default=str)
        
        # Delete original JSON from JSON_FILES_FOLDER
        original_json = review_session.get('json_files', {}).get(doc['name'])
        if original_json and os.path.exists(original_json):
            try:
                os.remove(original_json)
                print(f"[CLEANUP] Deleted: {os.path.basename(original_json)}")
            except Exception as e:
                print(f"[WARNING] Could not delete {original_json}: {e}")
        
        return jsonify({
            'success': True,
            'message': 'Document approved and saved',
            'outputPath': output_path,
            'filename': output_filename
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/review/document/<int:index>/reject', methods=['POST'])
def review_reject_document(index):
    """Reject document and save to REJECTED folder"""
    try:
        docs = review_session.get('documents', [])
        
        if index < 0 or index >= len(docs):
            return jsonify({'success': False, 'error': 'Invalid index'}), 400
        
        doc = docs[index]
        data = request.json
        fields = data.get('fields', review_session['extracted_data'].get(doc['name'], {}))
        reviewer = data.get('reviewer', 'Anonymous')
        notes = data.get('notes', '')
        reason = data.get('reason', 'Not specified')
        
        output_data = {
            'document_name': doc['name'],
            'document_type': doc['type'],
            'document_path': doc.get('path', ''),
            'document_location': doc.get('location', ''),
            'document_size': doc.get('size', 0),
            'review_status': 'REJECTED',
            'rejection_reason': reason,
            'reviewer': reviewer,
            'review_notes': notes,
            'review_timestamp': datetime.now().isoformat(),
            'processing_time': doc.get('processing_time', 0),
            'api_calls': doc.get('api_calls', {}),
            'extracted_fields': fields
        }
        
        base_name = Path(doc['name']).stem
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_filename = f"{base_name}_REJECTED_{timestamp}.json"
        output_path = os.path.join(REJECTED_FOLDER, output_filename)
        
        counter = 1
        while os.path.exists(output_path):
            output_filename = f"{base_name}_REJECTED_{timestamp}_{counter}.json"
            output_path = os.path.join(REJECTED_FOLDER, output_filename)
            counter += 1
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(output_data, f, indent=2, ensure_ascii=False, default=str)
        
        # Delete original JSON from JSON_FILES_FOLDER
        original_json = review_session.get('json_files', {}).get(doc['name'])
        if original_json and os.path.exists(original_json):
            try:
                os.remove(original_json)
                print(f"[CLEANUP] Deleted: {os.path.basename(original_json)}")
            except Exception as e:
                print(f"[WARNING] Could not delete {original_json}: {e}")
        
        return jsonify({
            'success': True,
            'message': 'Document rejected and saved',
            'outputPath': output_path,
            'filename': output_filename
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/review/stats', methods=['GET'])
def review_get_stats():
    """Get review statistics"""
    try:
        accepted_count = len([f for f in os.listdir(ACCEPTED_FOLDER) if f.endswith('.json')])
        rejected_count = len([f for f in os.listdir(REJECTED_FOLDER) if f.endswith('.json')])
        
        return jsonify({
            'success': True,
            'stats': {
                'inSession': len(review_session.get('documents', [])),
                'accepted': accepted_count,
                'rejected': rejected_count
            }
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ============================================================================
# MAIN ENTRY POINT
# ============================================================================
if __name__ == '__main__':
    print("=" * 80)
    print("[UNIVERSAL DATA EXTRACTOR] - All-in-One Server v2.0")
    print("=" * 80)
    print(f"[SERVER] http://localhost:{Config.SERVER_PORT}")
    print(f"[MAIN]   http://localhost:{Config.SERVER_PORT}")
    print(f"[REVIEW] http://localhost:{Config.SERVER_PORT}/review")
    print()
    print("[ENDPOINTS]")
    print("   POST /api/process/parallel  - Parallel document extraction")
    print("   POST /api/process           - Sequential extraction")
    print("   POST /api/pdf/split         - Intelligent PDF processing (split multi-invoice, process single normally)")
    print("   GET  /review                - Document Review UI")
    print("   POST /api/review/upload     - Upload docs for review")
    print("   POST /api/review/document/<n>/extract  - Extract fields")
    print("   POST /api/review/document/<n>/approve  - Approve document")
    print("   POST /api/review/document/<n>/reject   - Reject document")
    print()
    print("[HEALTH & MONITORING]")
    print("   GET  /health                - Basic health check")
    print("   GET  /ready                 - Readiness check")
    print()
    print("[CONFIGURATION]")
    print(f"   Max Threads: {Config.MAX_THREADS}")
    print(f"   Max Processes: {Config.MAX_PROCESSES}")
    print(f"   CORS Origins: {Config.CORS_ORIGINS}")
    print()
    print("[PDF SPLITTING]")
    if pdf_splitter:
        print(f"   [OK] PDF Processor: ENABLED (Unified Module)")
        print(f"   Output Path: {Config.SPLIT_PDF_OUTPUT_PATH}")
        print(f"   Method: {Config.SPLIT_PDF_METHOD}")
        print(f"   Auto Extract: {Config.SPLIT_PDF_AUTO_EXTRACT}")
        print(f"   Naming: {Config.SPLIT_PDF_NAMING_PATTERN}")
        
        # Show available features
        deps = pdf_splitter.check_dependencies()
        if deps['features'].get('barcode_detection', False):
            print(f"   [OK] Barcode Detection: Available")
        else:
            print(f"   [!!] Barcode Detection: Not available (missing libraries)")
        
        if deps['features'].get('ai_enhancement', False):
            print(f"   [OK] AI-Powered Enhancement: Available")
        else:
            print(f"   [!!] AI-Powered Enhancement: Not available (OpenAI not configured)")
        
        if deps['missing']:
            print(f"   [!!] Missing Optional Libraries: {', '.join(deps['missing'])}")
            print(f"   [!!] Install with: pip install {' '.join(deps['missing'])}")
    else:
        print(f"   [!!] PDF Processor: DISABLED")
        try:
            from pdf import BarcodeVendorPDFSplitter
            temp_proc = BarcodeVendorPDFSplitter()
            deps = temp_proc.check_dependencies()
            if deps['missing']:
                print(f"   [!!] Missing Libraries: {', '.join(deps['missing'])}")
                print(f"   [!!] Install with: pip install {' '.join(deps['missing'])}")
        except ImportError:
            print(f"   [!!] Reason: pdf.py module not found")
        except Exception as e:
            print(f"   [!!] Reason: {str(e)}")
    print()
    print("[MODULES LOADED]")
    if parallel_processor:
        print(f"   [OK] parallel_document_processor.py")
    else:
        print(f"   [!!] parallel_document_processor.py NOT loaded")
    if extractor:
        print(f"   [OK] multi_file_extractor.py")
    else:
        print(f"   [!!] multi_file_extractor.py NOT loaded")
    if pdf_processor:
        print(f"   [OK] pdf.py (Barcode & Vendor Detection - Priority #1)")
    else:
        print(f"   [!!] pdf.py NOT loaded")
    print()
    print(f"[OUTPUT DIRECTORIES]")
    print(f"   JSON Files: {JSON_FILES_FOLDER}")
    print(f"   Accepted: {ACCEPTED_FOLDER}")
    print(f"   Rejected: {REJECTED_FOLDER}")
    print(f"   Split PDFs: {SPLIT_PDF_FOLDER}")
    print("=" * 80)
    print()
    
    logger.info(f"Starting server on {Config.SERVER_HOST}:{Config.SERVER_PORT}")
    app.run(debug=Config.DEBUG_MODE, port=Config.SERVER_PORT, host=Config.SERVER_HOST)